"""Build SYNAPSE-Chapter-3.docx from CHAPTER-3-REVISED.md, figures included.

    python rasterize.py      # once, to produce diagrams/print/*.png
    python build_docx.py

What it does beyond a plain Markdown dump:

* every `*[placeholder - insert diagrams/fig-*.svg]*` line is replaced by the
  rasterised figure, scaled as large as the page allows and captioned below;
* each figure and each oversized table is given the page ORIENTATION that makes
  it biggest - dense diagrams are unreadable squeezed into a 6-inch column, so
  they get their own landscape page and the flow returns to portrait after;
* pipe tables become real Word tables with computed column widths, a repeating
  shaded header row, and rows that do not split across a page break;
* an auto-updating table of contents field is inserted after the front note.
"""
from __future__ import annotations

import re
import sys
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import (WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING,
                            WD_TAB_ALIGNMENT)
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Emu, Inches, Pt, RGBColor

# --------------------------------------------------------------- configuration

HERE = Path(__file__).resolve().parent
SOURCE = HERE / "CHAPTER-3-REVISED.md"
FIGURES = HERE / "diagrams" / "print"
OUTPUT = HERE.parent / "SYNAPSE-Chapter-3.docx"

BODY_FONT = "Times New Roman"
MONO_FONT = "Consolas"
BODY_PT = 12
LINE_SPACING = 1.5          # thesis body spacing
JUSTIFY_BODY = True
FIRST_LINE_INDENT = 0.0     # inches; set to 0.5 if the template wants indents

PAGE_W, PAGE_H = 8.5, 11.0
M_TOP, M_BOTTOM, M_INNER, M_OUTER = 1.0, 1.0, 1.5, 1.0   # inner = binding edge

PORTRAIT_TEXT_W = PAGE_W - M_INNER - M_OUTER             # 6.0"
PORTRAIT_TEXT_H = PAGE_H - M_TOP - M_BOTTOM              # 9.0"
LANDSCAPE_TEXT_W = PAGE_H - M_INNER - M_OUTER            # 8.5"
LANDSCAPE_TEXT_H = PAGE_W - M_TOP - M_BOTTOM             # 6.5"

CAPTION_RESERVE = 0.55      # vertical inches kept free for the figure caption

# Dense diagrams need a full page of their own, so a figure left exactly where
# the Markdown mentions it would strand a half-empty page before it. Figures
# therefore float to the end of the section that discusses them, the way a
# LaTeX float does. Set False to pin every figure to its source position.
FLOAT_FIGURES = True

CODE_INSET = 0.28     # left indent plus the border padding of a code block
CODE_MIN_PT, CODE_MAX_PT = 7.0, 9.5

RULE_GREY = RGBColor(0x44, 0x44, 0x44)
HEADER_FILL = "E8E8E8"
CODE_FILL = "F4F4F4"
QUOTE_FILL = "F7F7F7"

# A table turns landscape only when portrait would squeeze a prose column
# (one holding cells longer than PROSE_CHARS) below MIN_PROSE_WIDTH inches --
# the point where text breaks to two or three words a line and stops reading as
# a sentence. Width alone is a bad test: a three-row table is fine wrapped, and
# giving it a landscape page of its own wastes the page.
PROSE_CHARS = 60
MIN_PROSE_WIDTH = 1.6
DICTIONARY_HEADER = ["Attribute Name", "Data Type", "Max Length", "Key Type", "Null"]

EMDASH = "—"
FIGURE_LINE = re.compile(
    r"^\*\[placeholder\s+" + EMDASH + r"\s+insert\s+`diagrams/([^`]+)\.svg`\]\*\s*$"
)
CAPTION_LINE = re.compile(r"^\*\*(Figure\s+3\.\d+\s+" + EMDASH + r"\s+.+?)\*\*\s*$")
TABLE_CAPTION_LINE = re.compile(r"^\*\*(Table\s+3\.\d+\s+" + EMDASH + r"\s+.+?)\*\*\s*$")
HEADING_LINE = re.compile(r"^(#{1,6})\s+(.*)$")
LIST_LINE = re.compile(r"^(\s*)(?:([-*])|(\d+)\.)\s+(.*)$")
INLINE = re.compile(r"(`[^`]+`|\*\*.+?\*\*|(?<![\w*])\*(?!\s)[^*]+?(?<!\s)\*(?![\w*]))")


# --------------------------------------------------------------- block parsing


@dataclass
class Block:
    kind: str
    text: str = ""
    level: int = 0
    lines: tuple = ()
    rows: tuple = ()
    marker: str = ""


def parse(md: str) -> list[Block]:
    lines = md.split("\n")
    blocks: list[Block] = []
    i, n = 0, len(lines)
    buf: list[str] = []

    def flush() -> None:
        if buf:
            blocks.append(Block("para", text=" ".join(s.strip() for s in buf)))
            buf.clear()

    while i < n:
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            flush()
            i += 1
            continue

        if stripped.startswith("```"):
            flush()
            i += 1
            code: list[str] = []
            while i < n and not lines[i].strip().startswith("```"):
                code.append(lines[i])
                i += 1
            i += 1
            blocks.append(Block("code", lines=tuple(code)))
            continue

        # A figure caption immediately followed by its placeholder line.
        cap = CAPTION_LINE.match(stripped)
        if cap and i + 1 < n and FIGURE_LINE.match(lines[i + 1].strip()):
            flush()
            stem = FIGURE_LINE.match(lines[i + 1].strip()).group(1)
            blocks.append(Block("figure", text=cap.group(1), marker=stem))
            i += 2
            continue

        lone = FIGURE_LINE.match(stripped)
        if lone:
            flush()
            blocks.append(Block("figure", text="", marker=lone.group(1)))
            i += 1
            continue

        if stripped == "---":
            flush()
            blocks.append(Block("hr"))
            i += 1
            continue

        head = HEADING_LINE.match(stripped)
        if head:
            flush()
            blocks.append(Block("heading", text=head.group(2).strip(),
                                level=len(head.group(1))))
            i += 1
            continue

        if stripped.startswith(">"):
            flush()
            quote: list[str] = []
            while i < n and lines[i].strip().startswith(">"):
                quote.append(re.sub(r"^\s*>\s?", "", lines[i]))
                i += 1
            paras, cur = [], []
            for q in quote:
                if q.strip():
                    cur.append(q.strip())
                elif cur:
                    paras.append(" ".join(cur))
                    cur = []
            if cur:
                paras.append(" ".join(cur))
            blocks.append(Block("quote", lines=tuple(paras)))
            continue

        if stripped.startswith("|"):
            flush()
            rows = []
            while i < n and lines[i].strip().startswith("|"):
                cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                if not all(re.fullmatch(r":?-{2,}:?", c) for c in cells if c):
                    rows.append(cells)
                i += 1
            if rows:
                blocks.append(Block("table", rows=tuple(rows)))
            continue

        item = LIST_LINE.match(line)
        if item:
            flush()
            marker = item.group(2) or (item.group(3) + ".")
            body = [item.group(4).strip()]
            i += 1
            while (i < n and lines[i].strip()
                   and not LIST_LINE.match(lines[i])
                   and not lines[i].strip().startswith(("|", "#", ">", "```", "---"))):
                body.append(lines[i].strip())
                i += 1
            blocks.append(Block("list", text=" ".join(body), marker=marker))
            continue

        buf.append(line)
        i += 1

    flush()
    return blocks


# ----------------------------------------------------------------- low-level XML


def _shade(element, fill: str) -> None:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    element.append(shd)


def shade_paragraph(paragraph, fill: str) -> None:
    _shade(paragraph._p.get_or_add_pPr(), fill)


def box_paragraph(paragraph, colour: str = "C9C9C9", size: int = 4) -> None:
    borders = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement("w:" + edge)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(size))
        el.set(qn("w:space"), "6")
        el.set(qn("w:color"), colour)
        borders.append(el)
    paragraph._p.get_or_add_pPr().append(borders)


def left_bar(paragraph, colour: str = "999999") -> None:
    borders = OxmlElement("w:pBdr")
    el = OxmlElement("w:left")
    el.set(qn("w:val"), "single")
    el.set(qn("w:sz"), "18")
    el.set(qn("w:space"), "10")
    el.set(qn("w:color"), colour)
    borders.append(el)
    paragraph._p.get_or_add_pPr().append(borders)


def keep_with_next(paragraph) -> None:
    paragraph._p.get_or_add_pPr().append(OxmlElement("w:keepNext"))


def keep_lines(paragraph) -> None:
    paragraph._p.get_or_add_pPr().append(OxmlElement("w:keepLines"))


def field(paragraph, instruction: str, placeholder: str) -> None:
    """Insert a Word field (used for the TOC and the page number)."""
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " " + instruction + " "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = placeholder
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instr, separate, text, end):
        run._r.append(node)


def cell_borders(table) -> None:
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement("w:" + edge)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "6")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "808080")
        borders.append(el)
    table._tbl.tblPr.append(borders)


def table_margins(table, left_pt=5, right_pt=5, vertical_pt=2) -> None:
    margins = OxmlElement("w:tblCellMar")
    for edge, value in (("top", vertical_pt), ("left", left_pt),
                        ("bottom", vertical_pt), ("right", right_pt)):
        el = OxmlElement("w:" + edge)
        el.set(qn("w:w"), str(int(value * 20)))
        el.set(qn("w:type"), "dxa")
        margins.append(el)
    table._tbl.tblPr.append(margins)


def fixed_layout(table) -> None:
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    table._tbl.tblPr.append(layout)


def mark_header_row(row) -> None:
    row._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))


def no_split(row) -> None:
    row._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))


# ------------------------------------------------------------------ inline runs


def set_emphasis(run, bold: bool, italic: bool) -> None:
    """Turn emphasis on, never off: `False` must defer to the paragraph style,
    otherwise direct formatting un-bolds every heading and caption."""
    if bold:
        run.bold = True
    if italic:
        run.italic = True


def add_inline(paragraph, text, size=None, bold=False, italic=False, colour=None):
    """Render `code`, **bold** and *italic* into runs on `paragraph`."""
    for token in INLINE.split(text):
        if not token:
            continue
        if token.startswith("`") and token.endswith("`") and len(token) > 1:
            run = paragraph.add_run(token[1:-1])
            run.font.name = MONO_FONT
            run.font.size = Pt((size.pt if size else BODY_PT) - 1.5)
            set_emphasis(run, bold, italic)
        elif token.startswith("**") and token.endswith("**") and len(token) > 3:
            add_inline(paragraph, token[2:-2], size=size, bold=True,
                       italic=italic, colour=colour)
            continue
        elif token.startswith("*") and token.endswith("*") and len(token) > 1:
            add_inline(paragraph, token[1:-1], size=size, bold=bold,
                       italic=True, colour=colour)
            continue
        else:
            run = paragraph.add_run(token)
            run.font.name = BODY_FONT
            if size:
                run.font.size = size
            set_emphasis(run, bold, italic)
        if colour is not None:
            run.font.color.rgb = colour


def plain(text: str) -> str:
    """Markdown stripped, for width measurement."""
    return re.sub(r"[*`]", "", text)


# ----------------------------------------------------------------------- styles


def build_styles(doc) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(BODY_PT)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    pf = normal.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = LINE_SPACING
    pf.space_after = Pt(10)
    pf.space_before = Pt(0)
    if JUSTIFY_BODY:
        pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if FIRST_LINE_INDENT:
        pf.first_line_indent = Inches(FIRST_LINE_INDENT)

    def make(name, size, bold=False, italic=False, align=None, before=0, after=6,
             spacing=1.0, indent=0.0, hanging=0.0, font=BODY_FONT, keep=False,
             outline=None):
        style = doc.styles.add_style(name, 1)   # 1 = paragraph style
        style.base_style = doc.styles["Normal"]
        style.font.name = font
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.italic = italic
        style.font.color.rgb = RGBColor(0, 0, 0)
        style._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), font)
        fmt = style.paragraph_format
        fmt.space_before = Pt(before)
        fmt.space_after = Pt(after)
        fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        fmt.line_spacing = spacing
        fmt.alignment = align if align is not None else WD_ALIGN_PARAGRAPH.LEFT
        fmt.keep_with_next = keep
        if indent:
            fmt.left_indent = Inches(indent)
        if hanging:
            fmt.first_line_indent = Inches(-hanging)
        if outline is not None:
            el = OxmlElement("w:outlineLvl")
            el.set(qn("w:val"), str(outline))
            style._element.get_or_add_pPr().append(el)
        return style

    justified = WD_ALIGN_PARAGRAPH.JUSTIFY if JUSTIFY_BODY else WD_ALIGN_PARAGRAPH.LEFT

    make("ChapterTitle", 16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
         after=6, outline=0)
    make("ChapterSubtitle", 14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=24)
    make("H2", 14, bold=True, before=18, after=8, keep=True, outline=1)
    make("H3", 12.5, bold=True, before=14, after=6, keep=True, outline=2)
    make("H4", 12, bold=True, italic=True, before=12, after=6, keep=True, outline=3)
    make("Body", BODY_PT, spacing=LINE_SPACING, after=10, align=justified)
    make("ListEntry", BODY_PT, spacing=LINE_SPACING, after=8, indent=0.4,
         hanging=0.28, align=justified)
    make("CodeBlock", 9.5, font=MONO_FONT, before=6, after=10, indent=0.1)
    make("QuoteBlock", 11, spacing=1.15, before=4, after=8, indent=0.35, align=justified)
    make("FigureImage", BODY_PT, align=WD_ALIGN_PARAGRAPH.CENTER, before=6, after=4,
         keep=True)
    make("FigCaption", 10.5, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=14)
    make("TabCaption", 11, bold=True, before=12, after=4, keep=True)
    make("TableHead", 9.5, bold=True, after=0)
    make("TableCell", 9.5, after=0)
    make("TOCHeading", 14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=14)


# ------------------------------------------------------------------- page setup


def configure(section, landscape: bool) -> None:
    if landscape:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = Inches(PAGE_H), Inches(PAGE_W)
    else:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width, section.page_height = Inches(PAGE_W), Inches(PAGE_H)
    section.top_margin = Inches(M_TOP)
    section.bottom_margin = Inches(M_BOTTOM)
    section.left_margin = Inches(M_INNER)
    section.right_margin = Inches(M_OUTER)
    section.header_distance = Inches(0.5)
    section.footer_distance = Inches(0.5)


def add_footer(section) -> None:
    section.footer.is_linked_to_previous = False
    paragraph = section.footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    field(paragraph, "PAGE", "1")
    for run in paragraph.runs:
        run.font.name = BODY_FONT
        run.font.size = Pt(11)


# -------------------------------------------------------------------- builder


class Builder:
    def __init__(self, doc):
        self.doc = doc
        self.landscape = False
        self.pending_break = False

    def orient(self, landscape: bool) -> None:
        if landscape == self.landscape:
            return
        section = self.doc.add_section(WD_SECTION.NEW_PAGE)
        configure(section, landscape)
        section.footer.is_linked_to_previous = True
        self.landscape = landscape
        # The section break already opened a page; a manual one would blank a page.
        self.pending_break = False

    def take_break(self, paragraph) -> None:
        if self.pending_break:
            paragraph.paragraph_format.page_break_before = True
            self.pending_break = False

    @property
    def text_width(self) -> float:
        return LANDSCAPE_TEXT_W if self.landscape else PORTRAIT_TEXT_W

    def para(self, text: str, style: str = "Body"):
        self.orient(False)
        paragraph = self.doc.add_paragraph(style=style)
        self.take_break(paragraph)
        add_inline(paragraph, text)
        return paragraph

    def heading(self, text: str, level: int) -> None:
        self.orient(False)
        if level == 1:
            style = "ChapterTitle" if text.upper().startswith("CHAPTER") else "ChapterSubtitle"
            paragraph = self.doc.add_paragraph(style=style)
            self.take_break(paragraph)
            add_inline(paragraph, text.upper())
            keep_with_next(paragraph)
            return
        paragraph = self.doc.add_paragraph(style={2: "H2", 3: "H3", 4: "H4"}.get(level, "H4"))
        self.take_break(paragraph)
        add_inline(paragraph, text)
        keep_lines(paragraph)

    def code(self, lines: tuple) -> None:
        self.orient(False)
        paragraph = self.doc.add_paragraph(style="CodeBlock")
        self.take_break(paragraph)
        # Consolas advances 0.55 em per glyph; shrink until the widest line fits
        # on one row, because a wrapped listing loses its column alignment.
        longest = max((len(line) for line in lines), default=1)
        usable = self.text_width - CODE_INSET
        size = max(CODE_MIN_PT, min(CODE_MAX_PT, usable * 72.0 / (0.55 * longest)))
        for index, line in enumerate(lines):
            if index:
                paragraph.add_run().add_break()
            run = paragraph.add_run(line)
            run.font.name = MONO_FONT
            run.font.size = Pt(round(size, 1))
        shade_paragraph(paragraph, CODE_FILL)
        box_paragraph(paragraph)
        keep_lines(paragraph)

    def quote(self, paragraphs: tuple) -> None:
        self.orient(False)
        for index, text in enumerate(paragraphs):
            paragraph = self.doc.add_paragraph(style="QuoteBlock")
            self.take_break(paragraph)
            add_inline(paragraph, text, size=Pt(11))
            shade_paragraph(paragraph, QUOTE_FILL)
            left_bar(paragraph)
            if index < len(paragraphs) - 1:
                keep_with_next(paragraph)

    def bullet(self, marker: str, text: str) -> None:
        self.orient(False)
        paragraph = self.doc.add_paragraph(style="ListEntry")
        self.take_break(paragraph)
        paragraph.paragraph_format.tab_stops.add_tab_stop(
            Inches(0.4), WD_TAB_ALIGNMENT.LEFT
        )
        lead = paragraph.add_run(("•" if marker in "-*" else marker) + "\t")
        lead.font.name = BODY_FONT
        lead.bold = marker not in "-*"
        add_inline(paragraph, text)

    def page_break(self) -> None:
        self.pending_break = True

    # -- figures ------------------------------------------------------------

    def figure(self, stem: str, caption: str) -> None:
        png = FIGURES / (stem + ".png")
        if not png.exists():
            sys.exit("missing raster for " + stem + " - run rasterize.py first")
        head = (HERE / "diagrams" / (stem + ".svg")).read_text(encoding="utf-8")[:1200]
        native_w = float(re.search(r'width="([\d.]+)"', head).group(1)) / 96.0
        native_h = float(re.search(r'height="([\d.]+)"', head).group(1)) / 96.0

        def fit(avail_w, avail_h):
            return min(avail_w / native_w, avail_h / native_h)

        portrait = fit(PORTRAIT_TEXT_W, PORTRAIT_TEXT_H - CAPTION_RESERVE)
        landscape = fit(LANDSCAPE_TEXT_W, LANDSCAPE_TEXT_H - CAPTION_RESERVE)
        use_landscape = landscape > portrait * 1.02
        scale = min(landscape if use_landscape else portrait, 1.0)

        self.orient(use_landscape)
        paragraph = self.doc.add_paragraph(style="FigureImage")
        self.take_break(paragraph)
        paragraph.add_run().add_picture(str(png),
                                        width=Emu(int(native_w * scale * 914400)))
        keep_with_next(paragraph)
        if caption:
            add_inline(self.doc.add_paragraph(style="FigCaption"), caption,
                       size=Pt(10.5), bold=True)

    # -- tables -------------------------------------------------------------

    @staticmethod
    def allocate(ideal, avail, floor=0.55):
        """Share `avail` inches out in proportion to the ideal widths."""
        total = sum(ideal)
        widths = [max(floor, avail * w / total) for w in ideal]
        overflow = sum(widths) - avail
        if overflow > 0.001:                      # trim the widest columns back
            flexible = [i for i, w in enumerate(widths) if w > floor + 0.01]
            pool = sum(widths[i] - floor for i in flexible)
            for i in flexible:
                widths[i] -= overflow * (widths[i] - floor) / pool
        return widths

    @staticmethod
    def measure(rows, font_pt):
        """Ideal (unwrapped, capped) column widths in inches."""
        char = font_pt * 0.50 / 72.0
        widths = []
        for c in range(len(rows[0])):
            longest = max(len(plain(r[c])) for r in rows)
            header = len(plain(rows[0][c]))
            widths.append(max(header + 2, min(longest, 70)) * char)
        return widths

    def table(self, rows: tuple, caption) -> None:
        columns = max(len(r) for r in rows)
        rows = tuple(list(r) + [""] * (columns - len(r)) for r in rows)
        is_dictionary = [plain(c) for c in rows[0]] == DICTIONARY_HEADER

        font_pt = 9.5 if columns <= 5 else 8.5
        ideal = self.measure(rows, font_pt)
        longest = [max(len(plain(r[c])) for r in rows) for c in range(columns)]
        in_portrait = self.allocate(ideal, PORTRAIT_TEXT_W)
        wants_landscape = (not is_dictionary) and any(
            in_portrait[c] < MIN_PROSE_WIDTH
            for c in range(columns) if longest[c] > PROSE_CHARS
        )
        self.orient(wants_landscape)

        if caption:
            cap = self.doc.add_paragraph(style="TabCaption")
            self.take_break(cap)
            add_inline(cap, caption, size=Pt(11), bold=True)
        elif self.pending_break:
            spacer = self.doc.add_paragraph()
            spacer.paragraph_format.space_after = Pt(0)
            spacer.paragraph_format.line_spacing = 1.0
            self.take_break(spacer)

        avail = self.text_width
        if is_dictionary:
            widths = [avail * s for s in (0.34, 0.18, 0.145, 0.245, 0.09)]
        elif wants_landscape:
            widths = self.allocate(ideal, avail)
        else:
            widths = in_portrait

        table = self.doc.add_table(rows=0, cols=columns)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.autofit = False
        fixed_layout(table)
        cell_borders(table)
        table_margins(table)

        for r, source in enumerate(rows):
            row = table.add_row()
            no_split(row)
            if r == 0:
                mark_header_row(row)
            for c, value in enumerate(source):
                cell = table.cell(r, c)
                cell.width = Inches(widths[c])
                paragraph = cell.paragraphs[0]
                paragraph.style = self.doc.styles["TableHead" if r == 0 else "TableCell"]
                paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                add_inline(paragraph, value, size=Pt(font_pt), bold=(r == 0))
                if r == 0:
                    _shade(cell._tc.get_or_add_tcPr(), HEADER_FILL)
        for c, width in enumerate(widths):
            table.columns[c].width = Inches(width)

        spacer = self.doc.add_paragraph()
        spacer.paragraph_format.space_before = Pt(0)
        spacer.paragraph_format.space_after = Pt(6)
        spacer.paragraph_format.line_spacing = 1.0

    # -- contents -----------------------------------------------------------

    def toc(self) -> None:
        self.orient(False)
        self.pending_break = True
        heading = self.doc.add_paragraph(style="TOCHeading")
        self.take_break(heading)
        heading.add_run("TABLE OF CONTENTS")
        body = self.doc.add_paragraph()
        body.paragraph_format.line_spacing = 1.15
        field(body, 'TOC \\o "1-3" \\h \\z \\u',
              "Right-click and choose Update Field to build the contents.")
        for run in body.runs:
            run.font.name = BODY_FONT
            run.font.size = Pt(11)
            run.font.italic = True
            run.font.color.rgb = RULE_GREY


# -------------------------------------------------------------------------- main


def main() -> None:
    if not FIGURES.exists():
        sys.exit("diagrams/print is missing - run `python rasterize.py` first")

    blocks = parse(SOURCE.read_text(encoding="utf-8"))

    doc = Document()
    build_styles(doc)
    configure(doc.sections[0], landscape=False)
    add_footer(doc.sections[0])

    builder = Builder(doc)
    first_rule_seen = False
    pending_caption = None
    floating: list[Block] = []

    def land_figures() -> None:
        while floating:
            waiting = floating.pop(0)
            builder.figure(waiting.marker, waiting.text)

    for index, block in enumerate(blocks):
        nxt = blocks[index + 1] if index + 1 < len(blocks) else None

        if block.kind == "figure" and FLOAT_FIGURES:
            floating.append(block)
            continue
        if block.kind in ("heading", "hr"):
            land_figures()

        if block.kind == "heading":
            builder.heading(block.text, block.level)

        elif block.kind == "para":
            table_cap = TABLE_CAPTION_LINE.match(block.text.strip())
            if table_cap and nxt and nxt.kind == "table":
                pending_caption = table_cap.group(1)
                continue
            paragraph = builder.para(block.text)
            if nxt and nxt.kind in ("table", "code", "figure"):
                keep_with_next(paragraph)

        elif block.kind == "list":
            builder.bullet(block.marker, block.text)

        elif block.kind == "code":
            builder.code(block.lines)

        elif block.kind == "quote":
            builder.quote(block.lines)

        elif block.kind == "figure":
            builder.figure(block.marker, block.text)

        elif block.kind == "table":
            builder.table(block.rows, pending_caption)
            pending_caption = None

        elif block.kind == "hr":
            if not first_rule_seen:
                first_rule_seen = True
                builder.toc()
            builder.orient(False)
            builder.page_break()   # deferred: lands on the next block

    land_figures()

    doc.core_properties.title = "SYNAPSE - Chapter III: Methodology"
    doc.core_properties.subject = (
        "SYNAPSE: A Multi-Model ML and LLM HR ERP for Predictive Analytics and Recruitment"
    )
    doc.core_properties.author = "Earl Kian A. Bancayrin"
    doc.save(OUTPUT)

    figures = sum(1 for b in blocks if b.kind == "figure")
    tables = sum(1 for b in blocks if b.kind == "table")
    landscape = sum(1 for s in doc.sections if s.orientation == WD_ORIENT.LANDSCAPE)
    print("wrote " + str(OUTPUT))
    print("  {} blocks | {} figures | {} tables | {} sections ({} landscape)".format(
        len(blocks), figures, tables, len(doc.sections), landscape))


if __name__ == "__main__":
    main()
