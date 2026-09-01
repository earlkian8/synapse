"""Render CHAPTER-3-REVISIONS.md into a formatted Word document.

Handles the markdown subset actually used in that file: headings, pipe tables,
fenced code blocks, blockquotes (including tables and lists nested inside them),
horizontal rules, and inline bold / italic / code / links.
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent  # repo root
SRC = ROOT / "CHAPTER-3-REVISIONS.md"
OUT = ROOT / "CHAPTER-3-REVISIONS.docx"

BODY_FONT = "Calibri"
SERIF_FONT = "Georgia"
MONO_FONT = "Consolas"

INK = RGBColor(0x1A, 0x1A, 0x1A)
MUTED = RGBColor(0x5A, 0x5A, 0x5A)
ACCENT = RGBColor(0x1F, 0x3A, 0x5F)
ACCENT2 = RGBColor(0x8A, 0x2B, 0x1F)

SHADE_CODE = "F2F2F2"
SHADE_QUOTE = "FAF8F4"
SHADE_HEADER = "1F3A5F"
SHADE_BAND = "F4F6F9"


# ----------------------------------------------------------------------------
# low-level docx helpers
# ----------------------------------------------------------------------------


def shade(el, fill: str) -> None:
    pr = el.get_or_add_pPr() if el.tag.endswith("}p") else el
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    pr.append(shd)


def shade_cell(cell, fill: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def para_borders(p, *, left=None, box=None) -> None:
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    sides = []
    if box:
        sides = [("top", box), ("left", box), ("bottom", box), ("right", box)]
    elif left:
        sides = [("left", left)]
    for name, spec in sides:
        el = OxmlElement(f"w:{name}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(spec[0]))
        el.set(qn("w:space"), str(spec[1]))
        el.set(qn("w:color"), spec[2])
        pbdr.append(el)
    pPr.append(pbdr)


def add_hyperlink(paragraph, url: str, text: str, *, font=MONO_FONT, size=8):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rf = OxmlElement("w:rFonts")
    rf.set(qn("w:ascii"), font)
    rf.set(qn("w:hAnsi"), font)
    rPr.append(rf)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(size * 2)))
    rPr.append(sz)
    col = OxmlElement("w:color")
    col.set(qn("w:val"), "0B5FA5")
    rPr.append(col)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    run.append(t)
    link.append(run)
    paragraph._p.append(link)
    return link


def keep_with_next(p) -> None:
    pPr = p._p.get_or_add_pPr()
    el = OxmlElement("w:keepNext")
    pPr.append(el)


# ----------------------------------------------------------------------------
# inline markdown
# ----------------------------------------------------------------------------

INLINE = re.compile(
    r"(\*\*.+?\*\*|(?<!\*)\*(?!\*).+?(?<!\*)\*(?!\*)|`[^`]+`|https?://\S+)"
)


def write_inline(p, text: str, *, font=BODY_FONT, size=10.5, color=INK, italic=False,
                 bold=False):
    for token in INLINE.split(text):
        if not token:
            continue
        if token.startswith("**") and token.endswith("**") and len(token) > 4:
            r = p.add_run(token[2:-2])
            r.bold = True
        elif token.startswith("`") and token.endswith("`") and len(token) > 2:
            r = p.add_run(token[1:-1])
            r.font.name = MONO_FONT
            r.font.size = Pt(size - 1.0)
            r.font.color.rgb = ACCENT2
            continue
        elif token.startswith("http"):
            trail = ""
            while token and token[-1] in ".,);":
                trail = token[-1] + trail
                token = token[:-1]
            add_hyperlink(p, token, token, size=max(7.5, size - 2.0))
            if trail:
                p.add_run(trail)
            continue
        elif (
            token.startswith("*")
            and token.endswith("*")
            and len(token) > 2
            and not token.startswith("**")
        ):
            r = p.add_run(token[1:-1])
            r.italic = True
        else:
            r = p.add_run(token)
        r.font.name = font
        r.font.size = Pt(size)
        r.font.color.rgb = color
        if italic:
            r.italic = True
        if bold:
            r.bold = True
    return p


# ----------------------------------------------------------------------------
# block parsing
# ----------------------------------------------------------------------------


def parse_blocks(lines):
    """Yield (kind, payload) blocks from a list of markdown lines."""
    i, n = 0, len(lines)
    while i < n:
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if stripped.startswith("```"):
            j = i + 1
            buf = []
            while j < n and not lines[j].strip().startswith("```"):
                buf.append(lines[j])
                j += 1
            yield ("code", buf)
            i = j + 1
            continue

        if re.fullmatch(r"-{3,}", stripped):
            yield ("hr", None)
            i += 1
            continue

        m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if m:
            yield ("heading", (len(m.group(1)), m.group(2).strip()))
            i += 1
            continue

        if stripped.startswith("|"):
            rows = []
            while i < n and lines[i].strip().startswith("|"):
                rows.append(lines[i].strip())
                i += 1
            yield ("table", rows)
            continue

        if stripped.startswith(">"):
            buf = []
            while i < n and (lines[i].strip().startswith(">") or
                             (lines[i].strip() == "" and i + 1 < n and
                              lines[i + 1].strip().startswith(">"))):
                raw = lines[i]
                if raw.strip().startswith(">"):
                    buf.append(re.sub(r"^\s*>\s?", "", raw))
                else:
                    buf.append("")
                i += 1
            yield ("quote", buf)
            continue

        li = re.match(r"^(\d+)\.\s+(.*)$", stripped)
        bl = re.match(r"^[-*]\s+(.*)$", stripped)
        if li or bl:
            marker = f"{li.group(1)}." if li else "•"
            buf = [(li.group(2) if li else bl.group(1))]
            i += 1
            while i < n:
                nxt = lines[i]
                t = nxt.strip()
                if (not t or t.startswith(("|", ">", "#", "```"))
                        or re.fullmatch(r"-{3,}", t)
                        or re.match(r"^(\d+\.|[-*])\s+", t)):
                    break
                if not nxt.startswith((" ", "	")):
                    break
                buf.append(t)
                i += 1
            yield ("listitem", (marker, " ".join(buf)))
            continue

        # plain paragraph: gather until blank / structural line
        buf = [stripped]
        i += 1
        while i < n:
            nxt = lines[i].strip()
            if (not nxt or nxt.startswith(("|", ">", "#", "```"))
                    or re.fullmatch(r"-{3,}", nxt)
                    or nxt.startswith("**")):
                break
            buf.append(nxt)
            i += 1
        yield ("paragraph", " ".join(buf))


def split_table(rows):
    def cells(r):
        r = r.strip()
        if r.startswith("|"):
            r = r[1:]
        if r.endswith("|"):
            r = r[:-1]
        return [c.strip() for c in r.split("|")]

    parsed = [cells(r) for r in rows]
    body = [r for r in parsed if not all(re.fullmatch(r":?-{2,}:?", c or "-") for c in r)]
    if not body:
        return [], []
    return body[0], body[1:]


# ----------------------------------------------------------------------------
# renderers
# ----------------------------------------------------------------------------


def render_table(doc, rows, *, indent=0.0, size=9.0):
    header, body = split_table(rows)
    if not header:
        return
    ncols = len(header)
    table = doc.add_table(rows=0, cols=ncols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False

    widths_src = [max(len(header[c]), *(len(r[c]) if c < len(r) else 0 for r in body))
                  if body else len(header[c]) for c in range(ncols)]
    widths_src = [max(4, min(w, 90)) for w in widths_src]
    total = sum(widths_src)
    avail = 6.5 - indent
    widths = [Inches(max(0.45, avail * w / total)) for w in widths_src]

    hdr = table.add_row()
    for c, txt in enumerate(header):
        cell = hdr.cells[c]
        cell.width = widths[c]
        shade_cell(cell, SHADE_HEADER)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(re.sub(r"\*\*(.+?)\*\*", r"\1", txt))
        r.bold = True
        r.font.name = BODY_FONT
        r.font.size = Pt(size)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for ri, row in enumerate(body):
        tr = table.add_row()
        for c in range(ncols):
            cell = tr.cells[c]
            cell.width = widths[c]
            if ri % 2 == 1:
                shade_cell(cell, SHADE_BAND)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            write_inline(p, row[c] if c < len(row) else "", size=size)

    if indent:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.paragraph_format.left_indent = Inches(0.05)
        tblPr = table._tbl.tblPr
        ind = OxmlElement("w:tblInd")
        ind.set(qn("w:w"), str(int(indent * 1440)))
        ind.set(qn("w:type"), "dxa")
        tblPr.append(ind)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def render_quote(doc, lines):
    """Blockquote content: paragraphs, nested tables, list items."""
    for kind, payload in parse_blocks(lines):
        if kind == "table":
            render_table(doc, payload, indent=0.30, size=8.5)
        elif kind == "heading":
            _lvl, htext = payload
            p = doc.add_paragraph()
            pf = p.paragraph_format
            pf.left_indent = Inches(0.30)
            pf.right_indent = Inches(0.10)
            pf.space_before = Pt(8)
            pf.space_after = Pt(3)
            pf.keep_with_next = True
            shade(p._p, SHADE_QUOTE)
            para_borders(p, left=(12, 6, "C9B58A"))
            r = p.add_run(re.sub(r"\*\*(.+?)\*\*", r"", htext))
            r.bold = True
            r.font.name = SERIF_FONT
            r.font.size = Pt(11)
            r.font.color.rgb = INK
        elif kind == "hr":
            continue
        elif kind in ("paragraph", "listitem"):
            if kind == "listitem":
                marker, body = payload
                text = (f"{marker} {body}" if marker != "•"
                        else f"* {body}")
            else:
                text = payload
            li = re.match(r"^(\d+)\.\s+(.*)$", text)
            bullet = re.match(r"^[*-]\s+(.*)$", text)
            p = doc.add_paragraph()
            pf = p.paragraph_format
            pf.left_indent = Inches(0.30)
            pf.right_indent = Inches(0.10)
            pf.space_before = Pt(2)
            pf.space_after = Pt(6)
            pf.line_spacing = 1.14
            shade(p._p, SHADE_QUOTE)
            para_borders(p, left=(12, 6, "C9B58A"))
            if li:
                pf.left_indent = Inches(0.55)
                pf.first_line_indent = Inches(-0.25)
                r = p.add_run(f"{li.group(1)}.  ")
                r.bold = True
                r.font.name = SERIF_FONT
                r.font.size = Pt(10)
                r.font.color.rgb = INK
                text = li.group(2)
            elif bullet:
                pf.left_indent = Inches(0.55)
                pf.first_line_indent = Inches(-0.20)
                r = p.add_run("•  ")
                r.font.name = SERIF_FONT
                r.font.size = Pt(10)
                r.font.color.rgb = INK
                text = bullet.group(1)
            write_inline(p, text, font=SERIF_FONT, size=10)
        elif kind == "code":
            render_code(doc, payload, indent=0.30)


def render_code(doc, buf, *, indent=0.0, accent=False):
    text = "\n".join(buf).rstrip()
    if not text:
        return
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Inches(0.25 + indent)
    pf.right_indent = Inches(0.10)
    pf.space_before = Pt(4)
    pf.space_after = Pt(8)
    pf.line_spacing = 1.0
    shade(p._p, "FFF4D6" if accent else SHADE_CODE)
    para_borders(p, box=(6, 6, "D9C48F" if accent else "D0D0D0"))
    for k, ln in enumerate(text.split("\n")):
        if k:
            p.add_run().add_break()
        r = p.add_run(ln)
        r.font.name = MONO_FONT
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x22, 0x22, 0x22)


def render_heading(doc, level, text):
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    sizes = {1: 22, 2: 15, 3: 12, 4: 11}
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(2 if level == 1 else (18 if level == 2 else 12))
    pf.space_after = Pt(10 if level == 1 else 6)
    pf.keep_with_next = True
    r = p.add_run(text)
    r.bold = True
    r.font.name = BODY_FONT
    r.font.size = Pt(sizes.get(level, 10.5))
    r.font.color.rgb = ACCENT if level <= 2 else RGBColor(0x2E, 0x2E, 0x2E)
    if level == 2:
        para_borders(p, left=(18, 8, "1F3A5F"))
        pf.left_indent = Inches(0.06)
    return p



def build_contents(lines):
    """Derive the contents list from the markdown itself, so it cannot drift."""
    out, cur_ids = [], []
    pending = None

    def flush():
        if pending is not None:
            part, label = pending
            rng = None
            if cur_ids:
                lo, hi = cur_ids[0], cur_ids[-1]
                rng = ("R-%02d to R-%02d" % (lo, hi)) if hi - lo > 1 else (
                    "R-%02d, R-%02d" % (lo, hi) if hi != lo else "R-%02d" % lo)
            out.append((part, label, rng))

    for ln in lines:
        m2 = re.match(r"^## (Part [A-Z])\s*[—\-]\s*(.+?)\s*$", ln)
        m3 = re.match(r"^### (Section .+?)\s*$", ln)
        mr = re.match(r"^\*\*R-(\d+)\s", ln)
        if m2:
            flush(); cur_ids = []
            pending = (m2.group(1), m2.group(2))
        elif m3:
            flush(); cur_ids = []
            pending = ("", re.sub(r"\s+-\s+", " ", m3.group(1)))
        elif mr:
            cur_ids.append(int(mr.group(1)))
    flush()
    return out


def render_contents(doc, CONTENTS):
    render_heading(doc, 2, "Contents")
    for part, label, rng in CONTENTS:
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_before = Pt(1)
        pf.space_after = Pt(1)
        pf.left_indent = Inches(0.0 if part else 0.35)
        if part:
            r = p.add_run(part.ljust(8))
            r.bold = True
            r.font.name = BODY_FONT
            r.font.size = Pt(10)
            r.font.color.rgb = ACCENT
        r = p.add_run(label)
        r.font.name = BODY_FONT
        r.font.size = Pt(10 if part else 9.5)
        r.font.color.rgb = INK if part else MUTED
        if rng:
            r2 = p.add_run("   " + rng)
            r2.font.name = MONO_FONT
            r2.font.size = Pt(8.5)
            r2.font.color.rgb = ACCENT2
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


# ----------------------------------------------------------------------------
# document build
# ----------------------------------------------------------------------------


def build():
    md = SRC.read_text(encoding="utf-8")
    lines = md.split("\n")

    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.left_margin = Inches(1.0)
    sec.right_margin = Inches(1.0)
    sec.top_margin = Inches(0.9)
    sec.bottom_margin = Inches(0.9)

    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.14

    # page footer with page numbers
    footer_p = sec.footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_p.add_run()
    run.font.name = BODY_FONT
    run.font.size = Pt(8.5)
    run.font.color.rgb = MUTED
    for instr in ("begin", "PAGE", "end"):
        el = OxmlElement("w:fldChar") if instr in ("begin", "end") else OxmlElement("w:instrText")
        if instr in ("begin", "end"):
            el.set(qn("w:fldCharType"), instr)
        else:
            el.set(qn("xml:space"), "preserve")
            el.text = " PAGE "
        run._r.append(el)

    pending_label = None
    blocks = list(parse_blocks(lines))

    # inject a contents block immediately before the first "Part A" heading
    for k, (kind, payload) in enumerate(blocks):
        if kind == "heading" and payload[1].startswith("Part A"):
            blocks.insert(k, ("contents", None))
            break

    for idx, (kind, payload) in enumerate(blocks):
        if kind == "heading":
            level, text = payload
            render_heading(doc, level, text)
            continue

        if kind == "contents":
            render_contents(doc, build_contents(lines))
            continue

        if kind == "listitem":
            marker, text = payload
            p = doc.add_paragraph()
            pf = p.paragraph_format
            pf.left_indent = Inches(0.35)
            pf.first_line_indent = Inches(-0.25)
            pf.space_before = Pt(1)
            pf.space_after = Pt(4)
            r = p.add_run(marker.ljust(3) if marker != "•" else "•  ")
            r.bold = marker != "•"
            r.font.name = BODY_FONT
            r.font.size = Pt(10.5)
            r.font.color.rgb = ACCENT if marker != "•" else INK
            write_inline(p, text)
            continue

        if kind == "hr":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            para_borders(p, left=None, box=None)
            pPr = p._p.get_or_add_pPr()
            pbdr = OxmlElement("w:pBdr")
            bot = OxmlElement("w:bottom")
            bot.set(qn("w:val"), "single")
            bot.set(qn("w:sz"), "4")
            bot.set(qn("w:space"), "1")
            bot.set(qn("w:color"), "DDDDDD")
            pbdr.append(bot)
            pPr.append(pbdr)
            continue

        if kind == "table":
            render_table(doc, payload)
            continue

        if kind == "quote":
            render_quote(doc, payload)
            continue

        if kind == "code":
            render_code(doc, payload, accent=(pending_label == "search"))
            pending_label = None
            continue

        # paragraph
        text = payload

        # revision headers: **R-01 · 3.1 · ...**
        m = re.fullmatch(r"\*\*(R-\d+)\s*·\s*(.*?)\*\*", text)
        if m:
            p = doc.add_paragraph()
            pf = p.paragraph_format
            pf.space_before = Pt(14)
            pf.space_after = Pt(5)
            pf.keep_with_next = True
            r = p.add_run(m.group(1))
            r.bold = True
            r.font.name = BODY_FONT
            r.font.size = Pt(12)
            r.font.color.rgb = ACCENT2
            r2 = p.add_run("   " + m.group(2))
            r2.bold = True
            r2.font.name = BODY_FONT
            r2.font.size = Pt(10.5)
            r2.font.color.rgb = ACCENT
            continue

        if text.strip() == "sentence search:":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.keep_with_next = True
            r = p.add_run("SENTENCE SEARCH")
            r.bold = True
            r.font.name = BODY_FONT
            r.font.size = Pt(8.5)
            r.font.color.rgb = MUTED
            pending_label = "search"
            continue

        if text.startswith("revised sentence/paragraph"):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.keep_with_next = True
            r = p.add_run("REVISED SENTENCE / PARAGRAPH")
            r.bold = True
            r.font.name = BODY_FONT
            r.font.size = Pt(8.5)
            r.font.color.rgb = MUTED
            note = re.search(r"\*\((.+?)\)\*", text)
            if note:
                r2 = p.add_run("   " + note.group(1))
                r2.italic = True
                r2.font.name = BODY_FONT
                r2.font.size = Pt(8.5)
                r2.font.color.rgb = MUTED
            continue

        p = doc.add_paragraph()
        write_inline(p, text)
        if text.startswith("*") and text.endswith("*") and not text.startswith("**"):
            for r in p.runs:
                r.font.color.rgb = MUTED
                r.font.size = Pt(9.5)

    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    path = build()
    print("wrote", path)
